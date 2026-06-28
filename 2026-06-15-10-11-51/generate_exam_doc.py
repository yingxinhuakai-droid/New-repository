#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate 2026 EIA law exam Word document with answers and analysis"""

import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_question(doc, q, is_multi=False):
    # Question
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{q['id']}. {q['question']}")
    run.bold = True
    run.font.size = Pt(11)

    # Options
    for i, opt in enumerate(q['options']):
        p_opt = doc.add_paragraph()
        p_opt.paragraph_format.space_before = Pt(1)
        p_opt.paragraph_format.space_after = Pt(1)
        p_opt.paragraph_format.left_indent = Cm(0.8)
        letter = chr(65 + i)
        is_correct = letter in q['answer']
        run_opt = p_opt.add_run(f"({letter}) {opt}")
        run_opt.font.size = Pt(10.5)
        if is_correct:
            run_opt.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    # Answer
    p_ans = doc.add_paragraph()
    p_ans.paragraph_format.space_before = Pt(4)
    p_ans.paragraph_format.space_after = Pt(2)
    run_label = p_ans.add_run(u'\u3010\u7b54\u6848\u3011')
    run_label.bold = True
    run_label.font.size = Pt(10.5)
    run_label.font.color.rgb = RGBColor(0x1A, 0x6B, 0x3C)
    run_ans = p_ans.add_run(q['answer'])
    run_ans.bold = True
    run_ans.font.size = Pt(10.5)
    run_ans.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    # Analysis
    p_exp = doc.add_paragraph()
    p_exp.paragraph_format.space_before = Pt(2)
    p_exp.paragraph_format.space_after = Pt(6)
    run_exp_label = p_exp.add_run(u'\u3010\u89e3\u6790\u3011')
    run_exp_label.bold = True
    run_exp_label.font.size = Pt(10.5)
    run_exp_label.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    run_exp = p_exp.add_run(q['analysis'])
    run_exp.font.size = Pt(10.5)

def generate():
    # Load data
    with open('/Users/houjiasanshao/WorkBuddy/2026-06-15-10-11-51/exam_data_single.json', 'r', encoding='utf-8') as f:
        single_choices = json.load(f)
    with open('/Users/houjiasanshao/WorkBuddy/2026-06-15-10-11-51/exam_data_multi.json', 'r', encoding='utf-8') as f:
        multi_choices = json.load(f)

    doc = Document()

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(6)
    run_title = p_title.add_run(u'2026\u5e74\u73af\u8bc4\u5de5\u7a0b\u5e08\u8003\u8bd5\n\u6cd5\u89c4\u771f\u9898\u7b54\u6848\u53ca\u89e3\u6790')
    run_title.bold = True
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    # Disclaimer
    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_note.paragraph_format.space_after = Pt(16)
    run_note = p_note.add_run(u'\u26a0\ufe0f \u672c\u8d44\u6599\u7b54\u6848\u53ca\u89e3\u6790\u4ec5\u4f9b\u53c2\u8003\uff0c\u8bf7\u4ee5\u5b98\u65b9\u516c\u5e03\u4e3a\u51c6 \u26a0\ufe0f')
    run_note.bold = True
    run_note.font.size = Pt(12)
    run_note.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    # Separator
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run(u'\u2501' * 40)
    run_line.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run_line.font.size = Pt(10)

    # Part 1: Single Choice
    p_part1 = doc.add_paragraph()
    p_part1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_part1.paragraph_format.space_before = Pt(12)
    p_part1.paragraph_format.space_after = Pt(8)
    run_part1 = p_part1.add_run(u'\u7b2c\u4e00\u90e8\u5206  \u5355\u9879\u9009\u62e9\u9898\uff08\u517990\u9898\uff0c\u6bcf\u98981\u5206\uff09')
    run_part1.bold = True
    run_part1.font.size = Pt(14)
    run_part1.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    for q in single_choices:
        add_question(doc, q, is_multi=False)

    # Page break before part 2
    doc.add_page_break()

    # Separator
    p_line2 = doc.add_paragraph()
    p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line2 = p_line2.add_run(u'\u2501' * 40)
    run_line2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run_line2.font.size = Pt(10)

    # Part 2: Multiple Choice
    p_part2 = doc.add_paragraph()
    p_part2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_part2.paragraph_format.space_before = Pt(12)
    p_part2.paragraph_format.space_after = Pt(8)
    run_part2 = p_part2.add_run(u'\u7b2c\u4e8c\u90e8\u5206  \u4e0d\u5b9a\u9879\u9009\u62e9\u9898\uff08\u517430\u9898\uff0c\u6bcf\u98982\u5206\uff09')
    run_part2.bold = True
    run_part2.font.size = Pt(14)
    run_part2.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    p_part2_note = doc.add_paragraph()
    p_part2_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_part2_note.paragraph_format.space_after = Pt(8)
    run_part2_note = p_part2_note.add_run(u'\uff08\u6bcf\u9898\u7684\u5907\u9009\u9879\u4e2d\uff0c\u81f3\u5c11\u6709\u4e00\u4e2a\u7b26\u5408\u9898\u610f\uff0c\u591a\u9009\u3001\u5c11\u9009\u3001\u9519\u9009\u5747\u4e0d\u5f97\u5206\uff09')
    run_part2_note.font.size = Pt(10)
    run_part2_note.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for q in multi_choices:
        add_question(doc, q, is_multi=True)

    # Disclaimer page
    doc.add_page_break()

    p_end_line = doc.add_paragraph()
    p_end_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_end_line = p_end_line.add_run(u'\u2501' * 40)
    run_end_line.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run_end_line.font.size = Pt(10)

    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_end.paragraph_format.space_before = Pt(20)
    run_end = p_end.add_run(u'\u514d\u8d23\u58f0\u660e')
    run_end.bold = True
    run_end.font.size = Pt(14)
    run_end.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p_disclaimer = doc.add_paragraph()
    p_disclaimer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_disclaimer.paragraph_format.left_indent = Cm(2)
    p_disclaimer.paragraph_format.right_indent = Cm(2)
    run_disc = p_disclaimer.add_run(
        u'\u672c\u8d44\u6599\u6240\u63d0\u4f9b\u76842026\u5e74\u73af\u8bc4\u6cd5\u89c4\u771f\u9898\u7b54\u6848\u53ca\u89e3\u6790\uff0c'
        u'\u7cfb\u6839\u636e\u76f8\u5173\u6cd5\u5f8b\u6cd5\u89c4\u6761\u6587\u6574\u7406\u800c\u6210\uff0c'
        u'\u4ec5\u4f9b\u5b66\u4e60\u53c2\u8003\u4e4b\u7528\u3002\u7531\u4e8e\u6cd5\u5f8b\u6cd5\u89c4\u53ef\u80fd\u5b58\u5728\u4fee\u8ba2\u66f4\u65b0\uff0c'
        u'\u4e14\u7f16\u8005\u7406\u89e3\u53ef\u80fd\u5b58\u5728\u504f\u5dee\uff0c'
        u'\u7b54\u6848\u53ca\u89e3\u6790\u5185\u5bb9\u4e0d\u4fdd\u8bc1\u5b8c\u5168\u51c6\u786e\uff0c'
        u'\u8bf7\u4ee5\u5b98\u65b9\u516c\u5e03\u7684\u7b54\u6848\u548c\u6cd5\u5f8b\u6cd5\u89c4\u539f\u6587\u4e3a\u51c6\u3002'
        u'\u4f7f\u7528\u8005\u5e94\u81ea\u884c\u5224\u65ad\u5e76\u627f\u62c5\u4f7f\u7528\u672c\u8d44\u6599\u7684\u98ce\u9669\u3002'
    )
    run_disc.font.size = Pt(10)
    run_disc.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Save
    output_path = "/Users/houjiasanshao/WorkBuddy/2026-06-15-10-11-51/2026\u73af\u8bc4\u6cd5\u89c4\u771f\u9898\u7b54\u6848\u53ca\u89e3\u6790.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    generate()
